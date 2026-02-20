from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class DocumentChunk:
    text: str
    metadata: dict[str, str | int | float] = field(default_factory=dict)


@dataclass
class LoadedDocument:
    filename: str
    file_type: str
    text: str
    metadata: dict[str, str | int | float] = field(default_factory=dict)
    pages: list[str] = field(default_factory=list)


SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".txt", ".md", ".py", ".js", ".ts",
    ".tsx", ".jsx", ".json", ".yaml", ".yml", ".toml",
    ".csv", ".html", ".css", ".rs", ".go", ".java",
    ".c", ".cpp", ".h",
}
CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".tsx", ".jsx", ".json", ".yaml",
    ".yml", ".toml", ".css", ".rs", ".go", ".java",
    ".c", ".cpp", ".h", ".html",
}


def load_document(file_path: Path) -> LoadedDocument:
    ext = file_path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        return _load_pdf(file_path)
    elif ext == ".docx":
        return _load_docx(file_path)
    else:
        return _load_text(file_path)


def _load_pdf(file_path: Path) -> LoadedDocument:
    import fitz  # PyMuPDF

    doc = fitz.open(str(file_path))
    pages: list[str] = []
    full_text_parts: list[str] = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()  # type: ignore[union-attr]
        cleaned = _clean_text(text)
        if cleaned:
            pages.append(cleaned)
            full_text_parts.append(cleaned)

    doc.close()
    return LoadedDocument(
        filename=file_path.name,
        file_type="pdf",
        text="\n\n".join(full_text_parts),
        metadata={"page_count": len(pages), "source": str(file_path)},
        pages=pages,
    )


def _load_docx(file_path: Path) -> LoadedDocument:
    from docx import Document  # type: ignore[import-untyped]

    doc = Document(str(file_path))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = _clean_text(para.text)
        if text:
            paragraphs.append(text)

    full_text = "\n\n".join(paragraphs)
    return LoadedDocument(
        filename=file_path.name,
        file_type="docx",
        text=full_text,
        metadata={"paragraph_count": len(paragraphs), "source": str(file_path)},
    )


def _load_text(file_path: Path) -> LoadedDocument:
    ext = file_path.suffix.lower()
    file_type = "code" if ext in CODE_EXTENSIONS else "text"

    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    text = ""
    for encoding in encodings:
        try:
            text = file_path.read_text(encoding=encoding)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if not text:
        text = file_path.read_text(encoding="utf-8", errors="replace")

    return LoadedDocument(
        filename=file_path.name,
        file_type=file_type,
        text=text,
        metadata={"source": str(file_path), "language": ext.lstrip(".")},
    )


def _clean_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"[^\S\n]+", " ", text)
    return text
